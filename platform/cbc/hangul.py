"""Make a Korean .docx name a face that can actually draw its own text.

Word chooses a font per script, not per run: `w:rFonts/@ascii` is consulted for Latin and
`@eastAsia` for Hangul, and a run that names only the Latin slot leaves the other one to the
application, which picks whatever the machine happens to have. The document then opens in two
unrelated typefaces, and it opens differently on a different machine -- which is worse than
opening wrongly, because the author cannot see it.

Both Korean generators stamp all four slots as they write. The trouble with stamping at the
point of writing is that it has to be remembered at every one of forty call sites, and the
nine runs it is easiest to forget are the ones written by hand outside the run helpers: a
title page, a figure label, a tag appended to a bibliography entry. So the invariant is
enforced once, over the finished document, rather than asserted forty times:

    sweep(doc, "Apple SD Gothic Neo")

It stamps what was missed, and then it checks -- if anything is still unnamed after the sweep
it raises, because a silent repair that can itself fail silently is no better than the bug.
"""

from __future__ import annotations

import re

from docx.oxml.ns import qn

#: Hangul syllables plus the standalone jamo that appear in Korean typography.
HANGUL = re.compile(r"[가-힣ㄱ-ㅎㅏ-ㅣ]")

SLOTS = ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs")


def stamp(run, name: str):
    """Name `name` on every script slot of `run`, not only the Latin one."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for slot in SLOTS:
        rf.set(qn(slot), name)
    return run


def _paragraphs(doc):
    """Every paragraph in the document, including the ones inside tables and page furniture."""
    yield from doc.paragraphs
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for inner in cell.tables:
                    for r in inner.rows:
                        for c in r.cells:
                            yield from c.paragraphs
    for sec in doc.sections:
        for part in (sec.header, sec.footer, sec.first_page_header, sec.first_page_footer):
            yield from part.paragraphs


def _east_asia_face(run) -> str | None:
    rpr = run._element.rPr
    if rpr is None:
        return None
    rf = rpr.find(qn("w:rFonts"))
    return rf.get(qn("w:eastAsia")) if rf is not None else None


def sweep(doc, face: str) -> int:
    """Give every Hangul-bearing run an East Asian face, then prove none was left out.

    Returns the number of runs the sweep had to repair, which is worth printing: it is zero
    when the generator stamped everything itself, and a rising number is the sign that a new
    block of prose was written without going through the run helpers.
    """
    repaired = 0
    for par in _paragraphs(doc):
        for run in par.runs:
            if HANGUL.search(run.text) and not _east_asia_face(run):
                stamp(run, face)
                repaired += 1

    left = [r.text[:60] for p in _paragraphs(doc) for r in p.runs
            if HANGUL.search(r.text) and not _east_asia_face(r)]
    if left:
        raise AssertionError(
            f"{len(left)} Korean runs still name no East Asian face after the sweep; Word "
            f"would substitute a different typeface for each on every machine:\n  "
            + "\n  ".join(repr(x) for x in left[:5]))
    return repaired
