#!/usr/bin/env python3
"""Guard: the English and Korean editions of the report must describe the same work.

    ./.venv/bin/python platform/check_reports.py

Both editions unpack the same dict from cbc.report_data, so they cannot disagree about a
value they both read. What they CAN do is disagree about which values they mention: an edit
to one generator's prose that is not made in the other produces two documents that describe
different work under the same title.

This lives outside platform/tests/ on purpose. Both report generators RUN that suite, to take
the check count from its output rather than typing it into the prose, so a check placed there
that compares the two finished documents is circular: rebuilding one edition fails the suite
because the other has not been rebuilt yet, which stops the second rebuild from ever
happening. The dependency only works one way round, and this is that way round.

The two .docx files are not tracked (see .gitignore), so in a fresh clone there is nothing
to compare and this exits 0 saying so. That is a skip, not a pass, and it says which word it
means: a guard that reported success on an absent file would be worse than no guard. Build
both editions and it does the real comparison.
"""

from __future__ import annotations

import collections
import re
import zipfile

from pathlib import Path

REPO = Path(__file__).resolve().parents[1]
EN = REPO / "docs" / "CognitionBioChem_Report.docx"
KO = REPO / "docs" / "CognitionBioChem_Report_KO.docx"

PASS: list[str] = []
FAIL: list[str] = []


def check(name: str, ok: bool, detail: str = "") -> bool:
    (PASS if ok else FAIL).append(name)
    print(f"  {'PASS' if ok else 'FAIL'}  {name}" + (f"  -- {detail}" if detail else ""))
    return ok


def read(path: Path):
    from docx import Document
    d = Document(str(path))
    text = [x.text for x in d.paragraphs]
    for t in d.tables:
        for r in t.rows:
            text += [c.text for c in r.cells]
    media = len([n for n in zipfile.ZipFile(path).namelist()
                 if n.startswith("word/media/")])
    return d, text, media


_NUM = re.compile(r"-?\d[\d,]*\.?\d*")


def numbers(chunks) -> collections.Counter:
    c: collections.Counter = collections.Counter()
    for t in chunks:
        for x in _NUM.findall(t):
            c[x.rstrip(".").replace(",", "")] += 1
    return c


def main() -> int:
    print("=" * 76)
    print("CognitionBioChem report editions")
    print("=" * 76)
    print("\n[reports] the English and Korean editions agree")
    if not (EN.exists() and KO.exists()):
        built = [f.name for f in (EN, KO) if f.exists()] or ["neither"]
        print(f"  SKIP  nothing to compare -- the built editions are not tracked and are "
              f"not present here ({', '.join(built)} found)")
        print("        build both, then re-run:")
        print("          ./.venv/bin/python platform/build_report.py")
        print("          ./.venv/bin/python platform/build_report_ko.py")
        print("\n" + "=" * 76)
        print("0 passed, 0 failed")
        print("=" * 76)
        return 0
    check("both editions are present", True, f"{EN.name}, {KO.name}")
    den, en, en_media = read(EN)
    dko, ko, ko_media = read(KO)
    check("the two editions embed the same number of figures", en_media == ko_media,
          f"{en_media} vs {ko_media}")
    check("the two editions have the same number of tables",
          len(den.tables) == len(dko.tables), f"{len(den.tables)} vs {len(dko.tables)}")
    check("the two editions have the same number of paragraphs",
          len(den.paragraphs) == len(dko.paragraphs),
          f"{len(den.paragraphs)} vs {len(dko.paragraphs)}")

    # One-way, not an equality: Korean legitimately carries numerals English spells as words
    # -- 제3자 for "third-party", 세 차례 for "three times", 0을 지나 for "through zero".
    missing = numbers(en) - numbers(ko)
    check("every number in the English edition appears in the Korean one", not missing,
          f"missing from KO: {dict(list(missing.items())[:8])}")

    print("\n" + "=" * 76)
    print(f"{len(PASS)} passed, {len(FAIL)} failed")
    for f in FAIL:
        print("  -", f)
    print("=" * 76)
    return 1 if FAIL else 0


if __name__ == "__main__":
    raise SystemExit(main())
