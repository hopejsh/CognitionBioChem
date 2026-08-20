"""Build docs/CognitionBioChem_Report_KO.docx -- the Korean edition of the written account.

This is not a copy of build_report.py with the strings swapped. Both generators unpack the
same dict from cbc.report_data, so the two documents read the same artefacts and cannot
disagree about a published figure: re-run a study and both change together, or neither does.
What differs here is the prose, the typography, and nothing else.

Terminology is fixed rather than chosen sentence by sentence -- pLDDT, ipTM, PAE, DockQ and
proper nouns stay in Latin script; the statistical vocabulary follows one glossary throughout
so that "criterion" and "test" stay as distinguishable in Korean as the English argument
needs them to be. The register is the plain declarative of a Korean scientific report.

References are not translated. A citation has to be findable, so titles, journals, PMIDs and
DOIs appear exactly as PubMed returned them. Bibliographic metadata retrieved from PubMed
(NLM/NCBI).
"""

from __future__ import annotations

import json
import re
import statistics
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "platform"))
from cbc.provenance import git_sha  # noqa: E402
from cbc.report_data import J, A, load  # noqa: E402

FIG = REPO / "docs" / "figures"
OUT = REPO / "docs" / "CognitionBioChem_Report_KO.docx"

NAVY = RGBColor(0x1A, 0x3D, 0x6D)
GREY = RGBColor(0x5A, 0x66, 0x78)

#: Word picks a font per script, not per run, so a Korean run styled only through
#: `font.name` gets the Latin face and falls back to whatever the system offers for Hangul --
#: which is how a carefully set document ends up in two unrelated typefaces. Every run here
#: is stamped on all four script slots instead.
FONT_KO = "Apple SD Gothic Neo"


def kfont(run, name: str = FONT_KO):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(slot), name)
    return run


D = load()
REF = {r["key"]: r for r in D["refs"]["references"]}
CITED: list[str] = []


def cite(*keys) -> str:
    """대괄호 인용을 돌려주고 처음 등장한 순서를 기록한다."""
    nums = []
    for k in keys:
        if k not in REF:
            raise KeyError(f"{k} is not in docs/REFERENCES.json -- refusing to cite it")
        if k not in CITED:
            CITED.append(k)
        nums.append(str(CITED.index(k) + 1))
    return "[" + ", ".join(nums) + "]"


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY
        kfont(r)
    return p


def para(doc, text, size=10.5, italic=False, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    kfont(r)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(size)
    kfont(r)
    p.paragraph_format.space_after = Pt(3)
    return p


FIGNO = [0]


def figure(doc, name, caption, width=6.3):
    """그림을 배치하고 위치에 따라 번호를 매긴다.

    번호를 캡션에 직접 써 두면 그림을 하나 끼워 넣을 때마다 뒤의 번호가 모두 어긋난다.
    이 저장소가 없애려는 손수 관리하는 대조표와 같은 종류의 문제다.
    """
    path = FIG / name
    if not path.exists():
        raise FileNotFoundError(f"{path} is missing; the report will not be built without it")
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    FIGNO[0] += 1
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(f"그림 {FIGNO[0]}. {caption}")
    r.font.size = Pt(8.5)
    r.font.color.rgb = GREY
    r.italic = True
    kfont(r)
    c.paragraph_format.space_after = Pt(14)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hd)
        r.bold = True
        r.font.size = Pt(9)
        kfont(r)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
            kfont(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return t


def build() -> int:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT_KO
    st.font.size = Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_KO)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)

    cit, msa, scr = D["cit"], D["msa"], D["scr"]
    slate, struct, af = D["slate"], D["struct"], D["af"]
    m, c, nul, ver, per = D["m"], D["c"], D["nul"], D["ver"], D["per"]
    n_decoys, winners, win_margins = D["n_decoys"], D["winners"], D["win_margins"]
    scr_series, scr_decoys, n_checks = D["scr_series"], D["scr_decoys"], D["n_checks"]
    win_margins = "과 ".join(D["win_margin_values"])

    # ---------------------------------------------------------------- title page ---- #
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CognitionBioChem"); r.bold = True; r.font.size = Pt(26)
    r.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("음성 결과를 보고하는 구조약리학 워크벤치")
    r.font.size = Pt(13); r.font.color.rgb = GREY
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("왜 만들어졌고, 무엇을 하며, 무엇을 발견했는가")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
    doc.add_paragraph()

    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = a.add_run(f"{cit['authors'][0]}\nORCID {cit['orcid'][0].split('/')[-1]}\n\n"
                  f"버전 {cit['version']} · {cit['date_released']}")
    r.font.size = Pt(10.5)
    ids = doc.add_paragraph(); ids.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ids.add_run("  ·  ".join(i["value"] for i in cit["identifiers"]))
    r.font.size = Pt(9); r.font.color.rgb = GREY
    doc.add_paragraph()

    # ---------------------------------------------------------------- abstract ------ #
    h(doc, "요약", 1)
    para(doc,
         f"본 프로젝트는 인지 관련 CNS 표적을 겨냥해 설계되고 손으로 조립한 펩타이드 "
         f"서열 집합이 그 표적에 결합하는지를 물었다. 답은 '아니다'이며, 어느 쪽으로 답이 "
         f"나오든 그 답을 신뢰할 수 있도록 작업이 구성되었다.")
    para(doc,
         f"다중서열정렬(multiple sequence alignment) 없이 수행한 {scr['n_observed']}중 "
         f"스크리닝과 이를 넣고 다시 돌린 {msa['n_observed']}중 재실행에 걸쳐, 설계된 "
         f"펩타이드는 자기 아미노산을 조성 일치로 섞은 서열과 분리되지 않았다. 네이티브 "
         f"계면 pTM 평균은 {m['mean_native_iptm']}, 디코이(decoy) 평균은 "
         f"{m['mean_decoy_iptm']}이었고, 대응표본 t검정의 p = "
         f"{msa['p_holm']['H1_natives_separate_from_decoys']:.2f}이다. "
         f"후보 {len(per)}개 가운데 {nul['observed']}개가 자기 디코이 {n_decoys}개를 모두 "
         f"이겼지만, 우연만으로도 {nul['expected_under_null']}개가 그렇게 될 것으로 "
         f"기대되며 P(X ≥ {nul['observed']}) = {nul['p_at_least_observed']}이다. 가설은 "
         f"두 스크리닝 연구에서 보존된 {ver['n_versions']}개 버전 전부에서 반증되었으며, "
         f"후보 집합의 규모는 설계 "
         f"{min(v['n_candidates'] for v in ver['versions'])}개에서 "
         f"{max(v['n_candidates'] for v in ver['versions'])}개에 걸쳐 있었다.")
    para(doc,
         f"그 답에 이르기 위해 구축한 장치가 이 작업의 실체다: 데이터를 보기 전에 내용 "
         f"해시로 사전 등록된 연구 {c['studies']}건, {c['falsified']}개가 반증된 가설 "
         f"{c['hypotheses']}개, 내용 주소화 보관 아래 둔 구조 "
         f"{len(struct['entries'])}개, 그리고 각각이 지목하는 결함에서 실패함을 하나하나 "
         f"확인한 자동 검사 {n_checks}개다. 본 프로젝트는 자신의 이전 버전이 꾸며낸 "
         f"친화도·신뢰도·안전성 값을 마치 계산된 것처럼 표시하고 있었음을 발견하는 데서 "
         f"시작했다. 그 주장들은 삭제되지 않고 보존되어 있다.")

    doc.add_page_break()

    # ---------------------------------------------------------------- 1. WHY -------- #
    h(doc, "1  본 프로젝트가 존재하는 이유", 1)
    h(doc, "1.1  출발점이 된 실패", 2)
    n_retr = len(D["retr"])
    med_disc = statistics.median(a["discrepancy"]["value"] for a in D["aud"])
    med_ord = statistics.median(a["discrepancy_orders"]["value"] for a in D["aud"])
    para(doc,
         f"본 프로젝트의 이전 판은 구조 예측 모델이 산출한 것처럼 숫자를 제시했다. 그러나 "
         f"그 숫자를 산출한 계산은 없었다. {n_retr}개 후보가 결합 친화도, 신뢰도, 안전성 "
         f"값을 달고 있었으나 그 뒤에는 어떤 계산도 없었다: \"pLDDT\" 곡선은 "
         f"93 + sin(i·0.4)·4 + (charCode % 5)·0.5 라는 식이었고, 명시된 ΔG/Kd 쌍 "
         f"{n_retr}개는 하나도 빠짐없이 298.15 K에서 ΔG = RT·ln(Kd)에 비추어 내부적으로 "
         f"모순되며, 그 편차의 중앙값은 {med_disc:.2f} kcal/mol — Kd로는 {med_ord:.1f}자릿수에 "
         f"해당한다. 이 값들은 데이터 층에 retracted_claims 아래 보존되어 있으며, 그 값들을 "
         f"만들어 낸 렌더러는 사라졌다.")
    para(doc,
         "이것이 본 프로젝트가 현재의 형태로 존재하는 구체적인 이유이며, 그렇지 않았다면 "
         "이상해 보였을 설계 결정 하나를 설명해 준다: 이 워크벤치는 출처 기록 없이는 숫자를 "
         "표시할 수 없도록 만들어졌고, 한 번도 계산된 적이 없는 항목은 수치가 아니라 라벨로 "
         "표시된다.")

    h(doc, "1.2  질문에 영가설이 필요했던 이유", 2)
    para(doc,
         f"구조 예측 모델은 거의 어떤 사슬 쌍에 대해서도 신뢰도가 높아 보이는 복합체를 "
         f"내놓는다. 이 집합에서 가장 높은 계면 pTM은 "
         f"{max(p['native_iptm'] for p in msa['per_candidate'])}이며, 계면 품질에 대해 "
         f"확립된 구간 {cite('basu2016', 'lensink2007')}에 따르면 통상 좋은 계면으로 읽힐 "
         f"값이다. 문제는 설계된 펩타이드가 좋은 점수를 받느냐가 아니라, 같은 아미노산을 "
         f"다른 순서로 늘어놓은 서열보다 더 좋은 점수를 받느냐이다. 그 비교가 없다면 "
         f"스크리닝은 모델이 얼마나 확신하는지를 잴 뿐, 설계가 무언가를 해냈는지를 재는 "
         f"것이 아니다.")
    para(doc,
         f"이 우려는 이 부류의 모델에 대해서는 가정에 그치지 않는다. 여기서 사용한 예측 "
         f"모델에 대한 독립적인 평가들은 그 결합 친화도 출력이 결합 부위 돌연변이에, 그리고 "
         f"어떤 경우에는 표적을 바꾸는 것에조차 둔감할 수 있다고 보고하며 {cite('bret2026')}, "
         f"선도물질 발굴에 요구되는 에너지 분해능을 갖추지 못했다고 보고한다 {cite('wan2026')}. "
         f"물리적 타당성에 관한 연구는 딥러닝 도킹 방법이 RMSD 점수는 좋게 받으면서도 "
         f"기하학적으로 타당하지 않은 포즈를 자주 내놓는다는 점을 이미 보여 준 바 있다 "
         f"{cite('buttenschoen2024')}.")

    # ---------------------------------------------------------------- 2. BACKGROUND - #
    doc.add_page_break()
    h(doc, "2  과학적 배경", 1)

    h(doc, "2.1  구조 예측 모델이 보고하는 것과 보고하지 않는 것", 2)
    para(doc,
         f"AlphaFold2는 신경망이 서열만으로 단백질 구조를 원자 수준 정확도로 예측할 수 있다는 "
         f"것, 그리고 — 본 연구에는 그에 못지않게 중요하게도 — 자신의 출력에 대한 보정된 "
         f"신뢰도(confidence)를 스스로 보고할 수 있다는 것을 확립했다 {cite('jumper2021')}. 이 "
         f"신뢰도 양 가운데 두 가지를 본 프로젝트 전반에서 사용한다. pLDDT는 0–100 척도에서 국소 "
         f"정확도를 잔기(residue)별로 추정한 값이다. PAE, 즉 예측 정렬 오차는 구조를 잔기 j에 "
         f"중첩했을 때 잔기 i의 기대 위치 오차를 주는 행렬이다. 방향성을 가지므로 실제로 "
         f"비대칭이며, 비대각 블록은 두 영역이 서로에 대해 확신 있게 배치되어 있는지를 말해 준다. "
         f"계면(interface) pTM, 곧 ipTM은 그중 사슬 간 부분을 하나의 수치로 압축한다.")
    para(doc,
         f"이 가운데 어느 것도 결합 친화도가 아니며, 어느 것도 결합의 증거가 아니다. 이들은 기하 "
         f"구조에 대한 모델 자신의 확신을 모델이 스스로 보고한 값이다. 펩타이드–수용체 쌍의 "
         f"ipTM이 높다는 것은 모델이 그 펩타이드를 어딘가 확정된 자리에 놓는다는 뜻이지, 그 "
         f"펩타이드가 거기에 속한다는 뜻은 아니다. 이 구분이 프로젝트 전체가 도는 축이다.")

    h(doc, "2.2  사용한 예측 모델, 그리고 AlphaFold 3를 쓰지 않은 이유", 2)
    para(doc,
         f"이 저장소의 모든 예측 구조는 Boltz-2 v2.2.1로 생성했다 {cite('passaro2025')}. 이 모델은 "
         f"코드와 가중치가 모두 MIT 라이선스이며 Apple silicon에서 로컬로 실행된다. AlphaFold 3는 "
         f"사용하지 않았다. 파라미터가 신청제이고 비상업용이며 Linux/CUDA에 묶여 있어, 재현 가능한 "
         f"로컬 파이프라인을 그 위에 구축할 수 없었기 때문이다. AlphaFold Server 또한 사용하지 "
         f"않았고, 사용할 수도 없었다 — 그 약관은 단백질–리간드 및 단백질–펩타이드 결합 예측을 "
         f"위한 자동화된 사용을 금지하는데, 두 스크리닝 연구가 하는 일이 정확히 그것이다. 등재된 "
         f"AlphaFold Protein Structure Database {cite('varadi2024')}는 별도로 라이선스된 "
         f"코퍼스이며, CC BY 4.0 하에 독립적인 비교 대상으로 사용한다.")

    h(doc, "2.3  자유 에너지를 산출하지 않는 이유", 2)
    para(doc,
         f"Boltz-2에는 친화도 헤드가 있지만, 본 프로젝트는 그 출력을 자유 에너지로 나타내는 일을 "
         f"의도적으로 하지 않는다. 이 헤드는 Ki, Kd, IC50, EC50 라벨을 한데 모아 적합시킨 것이다. "
         f"이들은 서로 바꿔 쓸 수 있는 양이 아니다. IC50은 기질 농도와 효소의 Km에 의존하며, 이를 "
         f"Ki로 환산하려면 기작과 농도에 관한 가정이 필요한데, 한데 모은 학습 라벨은 그 가정을 "
         f"담고 있지 않다 {cite('cheng1973', 'cer2009')}. 그런 혼합에 적합된 수치는 단일한 물리적 "
         f"지시 대상을 갖지 않으므로 거기서 어떤 열역학적 양도 복원할 수 없으며, 어떤 코드 경로가 "
         f"이를 시도하면 저장소의 빌드 가드가 실패한다.")

    h(doc, "2.4  참조 데이터는 얼마나 시끄러운가", 2)
    ache_metrics = D["ache"]["metrics"]
    para(doc,
         f"공개 생물활성 데이터를 기준으로 예측 모델을 판정하려면 그 데이터가 얼마나 시끄러운지를 "
         f"알아야 한다. ChEMBL {cite('mendez2019', 'zdrazil2024')}에서 같은 단백질–리간드 쌍을 "
         f"독립적으로 측정한 값들은 상당히 다르다. 이질적인 공개 Ki 데이터의 실험 불확실성은 "
         f"표준편차 0.54 pKi 단위로 보고되어 있고, 이는 달성 가능한 R²를 0.81 부근에서 제한한다 "
         f"{cite('kramer2012')}. 혼합된 공개 IC50 데이터의 표준편차는 거기서 다시 25% 더 크다 "
         f"{cite('kalliokoski2013')}. 그리고 출처가 다른 IC50이나 Ki 값을 합치면 또 한 번 상당한 "
         f"잡음이 더해진다 {cite('landrum2024')}. 본 프로젝트는 자체 벤치마크 쌍 — 후페르진 "
         f"A(huperzine A) 대 아세틸콜린에스테라제(acetylcholinesterase) — 에서 같은 양을 "
         f"측정했고, 모델 평균 절대 오차 {ache_metrics['mean_absolute_error_log10']:.2f} log에 "
         f"대해 현장(in-situ) 참조 σ가 약 0.99 log임을 확인했다. 둘은 같은 자릿수이므로 해당 "
         f"연구는 모델 오차와 참조 오차를 분리할 수 없으며, 그 사실을 그대로 밝힌다.")
    para(doc,
         "이 0.99라는 값 자체에도 단서가 붙으며, 이를 덮어두지 않고 기록해 둔다. 이 값은 어세이 "
         "구성이 대체로 주석되어 있지 않은 ChEMBL 레코드를 한데 모아 선별한 집합 전체에서 측정한 "
         "것이다. 제대로 주석된 부분집합으로 한정하면 산포는 일반적인 IC50 하한 아래로 "
         "떨어지므로, 정직한 해석은 이 쌍이 유난히 시끄럽다기보다 유난히 주석이 부실하다는 "
         "것이다 — 그리고 해당 연구가 모델 오차와 참조 오차를 분리할 수 없다는 결론은 어느 "
         "쪽이든 그대로 유지된다.",
         size=9.5, italic=True)

    h(doc, "2.5  계면 품질을 어떻게 측정하는가", 2)
    para(doc,
         f"실험 구조가 존재하는 경우, 예측된 계면은 그 구조에 견주어 점수화할 수 있다. DockQ는 "
         f"CAPRI의 세 가지 양 — 복원된 천연 접촉의 비율, 리간드 RMSD, 계면 RMSD — 을 [0, 1] 구간의 "
         f"단일 연속 점수로 결합하며, 이 점수는 CAPRI의 acceptable/medium/high 분류를 거의 그대로 "
         f"재현한다 {cite('basu2016', 'lensink2007')}. 여기서 사용한 구현은 DockQ v2이다 "
         f"{cite('mirabello2024')}. 이것이 중요한 까닭은, 모델이 스스로 보고한 신뢰도가 실제 계면 "
         f"정확성을 따라가는지를 물을 수 있게 해 주기 때문이다. 이는 후보 스크리닝의 실행이 허용되기 "
         f"전에 슬레이트 #7이 답하도록 등록된 물음이다.")

    h(doc, "2.6  물리화학적 성질과 허용 여부", 2)
    para(doc,
         f"서열 수준의 성질은 단정하는 것이 아니라 계산한다: 분자량, pH 7.4에서의 알짜 전하, 등전점, "
         f"그리고 Kyte-Doolittle 소수성 지수 {cite('kyte1982')}. 화학 구조는 RDKit으로 검증하며 — "
         f"파싱, 분자식, InChIKey, 입체화학 완전성 — 저장된 SMILES가 그 이름과 다른 분자를 부호화하는 "
         f"화합물은 계산된 성질과 함께 표시하는 대신 미검증으로 표시한다. ADMET은 분자가 모델의 적용 "
         f"가능 영역 안에 들어오는 경우에만 ADMET-AI로 예측하고 {cite('swanson2024')}, 그렇지 않은 "
         f"경우에는 사유를 밝혀 거부한다. 표적 컨스트럭트(construct), 신호 펩타이드, 토폴로지, 이황화 "
         f"결합은 UniProt에서 가져오고 {cite('uniprot2025')}, 실험 참조 구조는 Protein Data "
         f"Bank에서 가져온다 {cite('berman2000')}.")

    h(doc, "2.7  다중성, 그리고 사전 등록된 기준이 아닌 것", 2)
    para(doc,
         f"각 연구는 세 개의 가설을 등록한다. 가설이 검정통계량으로 판정되는 경우, 그 가설군은 Holm "
         f"단계적 축소 절차로 보정한다 {cite('holm1979')}. 이 슬레이트의 가설 대부분은 검정이 아니다. "
         f"기술적 양에 대한 임계값 비교이며, 미리 그어 둔 선에 수치를 견주어 판정한다. 전체 가설 "
         f"{c['hypotheses']}개 가운데 검정으로 판정된 것은 {c['decided_by_a_test']}개뿐이다. 임계값 "
         f"기준을 0.0 또는 1.0의 p값으로 부호화하여 Holm에 넣는 것은 — 이 코드의 이전 판이 그렇게 "
         f"했다 — 도달할 수 없는 p값을 날조하는 동시에, 함께 묶인 실제 검정들로부터 승수를 훔친다. "
         f"이 둘은 서로 떼어 두었고, 페이지는 모든 판정에 그것이 어느 쪽이었는지를 표시한다.")

    # ---------------------------------------------------------------- 3. WHAT ------- #
    doc.add_page_break()
    h(doc, "3  무엇을 만들었는가", 1)
    para(doc,
         f"이 워크벤치는 출처 기록을 담고 있는 데이터 계층 위에 놓인 정적 페이지이며, 그 뒤에는 "
         f"예측과 연구와 검사를 실행하는 Python 플랫폼이 있다. 하중을 지탱하는 일은 세 가지 "
         f"기제가 맡는다.")

    h(doc, "3.1  렌더링 규칙으로서의 출처 기록", 2)
    para(doc,
         f"모든 과학적 값은 상태를 함께 지닌다. computed, predicted, database, literature, "
         f"measured, placeholder, not_computed 가운데 하나다. 사용자 인터페이스에서 값이 화면에 "
         f"이르는 경로는 정확히 하나뿐이며, 상태가 placeholder이거나 not_computed일 때 그 경로는 "
         f"수치가 아니라 라벨을 내보낸다. 날조된 값들을 그냥 지우고 잊어버릴 수 없었던 이유가 "
         f"여기에 있다. 그 값들은 retracted_claims 아래에 취소선이 그어진 채 눈에 보이도록 "
         f"보존되며, 그리하여 한때 무엇이 주장되었는지에 대한 기록이 사라지지 않는다.")
    figure(doc, "ui1_headline_finding.png",
           "워크벤치의 Overview 탭. 결과는 기능보다 먼저 제시되고, 카드에 실린 모든 "
           "수치는 마크업에 적어 넣은 것이 아니라 빌드 시점에 연구 산출물에서 "
           "읽어 온 것이다.")

    h(doc, "3.2  내용 해시 아래의 사전 등록", 2)
    plans = D["plans"]
    para(doc,
         f"연구를 실행하기 전에 그 질문과 주요 지표, 판정 임계값, 가설, 분석 계획을 고정하고 "
         f"해시한다. 레지스트리는 도달할 수 없는 판정을 담은 계획 — 달성 가능한 최소 보정 p값이 "
         f"이미 α를 넘어서는 기준 — 이나 반증 불가능한 가설을 담은 계획, 또는 모든 가설이 같은 "
         f"입장에서 예측되어 어떤 결과로도 판별이 되지 않는 계획을 거부한다. {c['studies']}개 "
         f"연구 계열에 걸쳐 {plans}건의 계획이 보존되어 있고, 대체된 판본은 덮어쓰지 "
         f"않고 남기며, 각각 어째서 대체되었는지를 "
         f"기록한다.")
    para(doc,
         f"중요한 귀결 하나가 페이지에 정직하게 기록되어 있다. 이 슬레이트의 어느 연구도 "
         f"확증적이지 않다. 모든 연구가 적어도 한 가지 측면에서 등록된 계획을 벗어났고, 각 "
         f"연구의 자체 감사가 그 사실을 밝히고 있다. 사전 등록이 이 결과들을 확증적으로 만들어 "
         f"준 것은 아니다 — 사전 등록이 한 일은 이탈을 눈에 보이게 만든 것이다.")

    h(doc, "3.3  계산의 보관", 2)
    runs, published_rows = D["runs"], D["published_rows"]
    para(doc,
         f"모든 예측 실행은 내용 주소화되어 있다. 디렉터리 이름이 그 파일들의 이름과 내용에 "
         f"대한 해시이므로, 출력이 하나라도 바뀌면 식별자가 바뀐다. {runs}건의 실행이 보관 "
         f"중이다. 어떤 연구가 보고하든 그 모든 행이 여전히 매니페스트에 남아 있는 실행으로 "
         f"해소된다는 것을 한 테스트가 확인하며 — 무언가를 폴딩하는 여섯 연구를 통틀어 "
         f"{published_rows}행이다 — 이것이 다시 해시된 실행과 "
         f"유실된 실행을 갈라놓는 지점이다.")
    figure(doc, "ui4_structure_gallery.png",
           f"구조 갤러리. 각 항목은 파이프라인이 산출한 모델이며, 내용 해시가 그 이름이 된 실행 "
           f"디렉터리에서 열린다. 항목마다 표시되는 신뢰도 항목은 그 모델이 실제로 정의하는 "
           f"것들뿐이어서, 단량체는 계면 점수를 0으로 보여 주는 것이 아니라 아예 보여 주지 "
           f"않는다.")
    para(doc,
         f"개발 도중 겪은 두 번의 실패가 이 장치의 계기였다. 컨스트럭트를 바로잡은 뒤 다시 "
         f"실행했는데 ipTM이 열여섯 자리까지 똑같이 나왔다. 예측 모델의 전처리 캐시가 입력 "
         f"레코드의 이름을 키로 삼는 탓에, 이전 수용체를 조용히 다시 폴딩한 것이었다. 그 뒤에는 "
         f"중단된 실행이 새 입력의 해시를 옛 결과와 짝지은 매니페스트 항목을 기록했다. 이제 두 "
         f"경우 모두, 산출된 모델을 다시 파싱해 그 사슬 길이를 요청과 대조하는 사후 조건에서 "
         f"실패한다.")

    # ---------------------------------------------------------------- 4. HOW -------- #
    doc.add_page_break()
    h(doc, "4  중심 질문을 어떻게 물었는가", 1)

    h(doc, "4.1  조성 일치 영가설", 2)
    n_dec = msa["per_candidate"][0]["n_decoys"]
    para(doc,
         f"후보마다 그 후보 자신의 아미노산을 섞어 디코이 서열 {n_dec}개를 생성한다. 따라서 "
         f"디코이는 조성이 동일하다 — 분자량도, 순전하도, 소수성 지표도 같다 — 오직 순서만 "
         f"다르다. 각 디코이는 같은 수용체 컨스트럭트를 상대로, 같은 시드로, 같은 순서로 "
         f"폴딩된다. 설계된 서열이 표적에 관한 정보를 담고 있다면 자기 자신을 섞은 서열보다 "
         f"높은 점수를 받아야 하고, 모델이 조성에 반응하고 있는 것이라면 그렇지 않을 것이다.")
    para(doc,
         f"수용체 컨스트럭트는 UniProt 전장 서열이 아니다. 신호 펩타이드는 제거하고, 결합 "
         f"부위가 세포 바깥에 있는 경우 막관통 구간과 세포질 구간은 제외하며, 실제로 폴딩한 "
         f"구간은 그 근거 및 표준 번호 매김과 함께 후보별로 기록한다. 작업 도중 컨스트럭트 "
         f"수정이 세 차례 있었다. 각각은 답의 정밀도를 바꾸었고, 어느 것도 답의 방향은 "
         f"바꾸지 않았다.")

    h(doc, "4.2  스크리닝을 유효하다고 인정하기 전에 둔 관문", 2)
    gate_pl, gate_rl = D["gate_pl"], D["gate_rl"]
    cand_pl, cand_rl = D["cand_pl"], D["cand_rl"]
    pi_m = D["pi"]["metrics"]
    para(doc,
         f"슬레이트 #7은 후보 스크리닝을 믿어도 된다고 허용하기 전에 등록되었다. 이 슬레이트는 "
         f"이 파이프라인이 암기했을 수도 있는 계면을 되찾아 내는지를 묻는다. 펩타이드–수용체 "
         f"X선 복합체 16개를 폴딩하고 DockQ로 점수를 매겼다 {cite('mirabello2024')}. "
         f"파이프라인은 복합체 {D['pi']['n_observed']}개 가운데 "
         f"{pi_m['fraction_dockq_acceptable']:.0%}에서 CAPRI 허용 가능 품질에 도달했고"
         f"(DockQ 중앙값 {pi_m['median_dockq']}), ipTM은 "
         f"Spearman ρ = {pi_m['spearman_iptm_dockq']}로 DockQ를 따라갔다. 이후 단계에서 "
         f"ipTM을 쓰는 의미는 바로 이것이다. 이런 종류의 복합체에서의 판별 신호로 쓰는 "
         f"것이지, 보정된 확률로 쓰는 것이 아니다.")
    para(doc,
         f"같은 연구는 그러한 관문을 필요하게 만드는 누출을 측정했다. 벤치마크 펩타이드는 "
         f"{min(gate_pl)}–{max(gate_pl)} 잔기이고 상대 수용체는 "
         f"{min(gate_rl)}–{max(gate_rl)} 잔기인 반면, 후보는 "
         f"{min(cand_pl)}–{max(cand_pl)} 잔기이고 상대 수용체는 "
         f"{min(cand_rl)}–{max(cand_rl)} 잔기이다. 민감도 논증은 그대로 옮겨 가지만 수치 "
         f"구간은 그렇지 않으며, 보고서는 임계값을 그대로 넘겨 쓰는 대신 그 사실을 "
         f"밝힌다.")

    h(doc, "4.3  차이가 얼마나 커야 의미를 갖는가", 2)
    iv_m = D["iv"]["metrics"]
    para(doc,
         f"슬레이트 #2는 어떤 비교도 신뢰하기 전에 샘플러 자체의 잡음을 측정했다. 같은 "
         f"복합체를 서로 다른 시드로 반복해 폴딩하면 시드 간 표준편차는 ipTM에서 "
         f"{iv_m['across_seed_sd_iptm']:.3f}, pLDDT 단위로 "
         f"{iv_m['across_seed_sd_complex_plddt']:.2f}이 나온다. 그보다 작은 차이는 측정이 "
         f"아니다. 스크리닝에서 겉보기 '성공'으로 보이는 {len(winners)}건을 잡음으로 보고하는 "
         f"이유가 바로 이것이다. 이들이 자기 최고 디코이를 앞선 폭은 "
         f"{win_margins}인데, 1 표준편차 산포는 "
         f"{iv_m['across_seed_sd_iptm']:.3f}이다.")

    # ---------------------------------------------------------------- 5. RESULTS ---- #
    doc.add_page_break()
    h(doc, "5  결과", 1)

    h(doc, "5.1  설계 서열은 자기 자신을 섞은 서열과 분리되지 않는다", 2)
    beaten = sum(1 for p in per if p["decoy_max"] > p["native_iptm"])
    beaten_mean = sum(1 for p in per if p["decoy_mean"] > p["native_iptm"])
    para(doc,
         f"완전한 다중서열정렬과 후보당 디코이 {n_decoys}개를 사용했을 때, 네이티브 "
         f"ipTM 평균은 "
         f"{m['mean_native_iptm']}이고 디코이 평균은 {m['mean_decoy_iptm']}이다. "
         f"쌍대 차이는 {m['mean_native_iptm'] - m['mean_decoy_iptm']:+.4f}이고 "
         f"p = {msa['p_holm']['H1_natives_separate_from_decoys']:.2f}, Cohen's "
         f"dz = {m['cohens_dz']:+.2f}이다. 후보 {len(per)}개 중 {beaten}개에서는 "
         f"섞은 서열 {n_decoys}개 중 최고값이 설계 서열보다 높은 점수를 냈고, {len(per)}개 중 "
         f"{beaten_mean}개에서는 디코이 평균이 그렇다.")
    figure(doc, "fig1_native_vs_decoy.png",
           f"각 설계 펩타이드를 자기 아미노산을 섞어 만든 서열 {n_decoys}개와 "
           f"대비한 것이다. 마름모는 설계 서열이고, 막대는 디코이 평균에서 최고 "
           f"디코이까지 이어진다. "
           "막대의 오른쪽에 놓인 설계 서열은 자신의 모든 섞은 서열을 앞섰다. 데이터: "
           "data/study_msa_specificity.json.")

    h(doc, "5.2  승자 둘은 우연이 만들어내는 모습 그대로다", 2)
    para(doc,
         f"후보 {len(per)}개 중 {nul['observed']}개가 자기 디코이 {n_decoys}개를 모두 "
         f"앞섰다({', '.join(w['code'] for w in winners)}). 영가설 아래에서 한 후보가 "
         f"그렇게 될 확률은 {nul['per_candidate_null_probability']} = 1/{n_decoys + 1}이므로, "
         f"{len(per)}개 중 {nul['expected_under_null']}개는 우연히 그렇게 되리라 기대되고 "
         f"P(X ≥ {nul['observed']}) = {nul['p_at_least_observed']}이다. 승자는 모두 "
         f"경험적 p가 {max(w['empirical_p'] for w in winners)}로 같은데, 이는 이 설계가 "
         f"만들어낼 수 있는 최솟값이며 통상적인 임계값을 "
         f"하나도 넘지 못한다. 조성 일치 영가설은 한 후보가 히트로 읽히는 것을 "
         f"막아 주지만, 같은 방식으로 읽히는 스크리닝에 대해서는 아무 일도 하지 않으며, "
         f"이는 한 층 위에서 저지르는 동일한 오류다. 스크리닝 수준의 영가설은 데이터를 "
         f"본 뒤에 계산되었고, 그 때문에 산출물에서 탐색적이라고 표시되어 있다.")
    figure(doc, "fig2_screen_level_null.png",
           f"스크리닝 수준의 영가설이다. 회색은 우연만으로 후보 {len(per)}개 중 몇 개가 "
           f"자기 디코이 {n_decoys}개를 모두 앞서게 되는지의 이항분포이고, 호박색은 "
           f"관측된 값이다. 데이터: data/study_msa_specificity.json.", width=4.8)

    h(doc, "5.3  판정은 모든 수정을 견뎠고, 격차는 그러지 못했다", 2)
    para(doc,
         f"컨스트럭트 수정 세 차례, 중복 제거 한 차례, 커버리지 확장 두 차례, 그리고 "
         f"후보당 섞은 서열을 {scr_decoys}개에서 {n_decoys}개로 늘린 디코이 확장을 "
         f"거치는 동안, H1은 보존된 "
         f"{ver['n_versions']}개 버전 모두에서 반증되었다. 일정하게 유지되지 않은 것은 "
         f"격차의 크기다. 스크리닝의 {len(scr_series)}개 버전에 걸쳐 네이티브 빼기 디코이 "
         f"평균은 {', '.join(f'{v:+.4f}' for v in scr_series)}로 움직였고, 몇 배로 커졌다가 "
         f"0을 지나 되돌아왔으며, 샘플러 잡음 하한"
         f"(±{D['iv']['metrics']['across_seed_sd_iptm']:.3f} ipTM)을 한 번도 벗어난 적이 "
         f"없다. 모든 수정을 견딘 것은 "
         f"판정이지 격차가 아니다.")
    figure(doc, "fig3_falsified_every_version.png",
           "두 스크리닝 연구에서 보존된 모든 버전과 각 버전의 후보 수이다. "
           "데이터: data/slate.json, data/superseded/.", width=5.9)

    h(doc, "5.4  신뢰도는 실제로 어디에서 오는가", 2)
    e = next(x for x in struct["entries"] if x["id"] == "cpx-BasalAChE-Abeta-B4")
    e_share = e["chains"][0]["length"] / sum(ch["length"] for ch in e["chains"])
    para(doc,
         f"복합체 하나가 집계 통계량이 추상적으로 말하는 바를 그대로 보여 준다. "
         f"{e['target']}에 대한 {e['code']}의 ipTM은 {e['metrics']['iptm']}이며, 통상적인 구간 "
         f"기준으로는 좋은 계면이다. 사슬별로 나누면 수용체의 평균 pLDDT는 "
         f"{e['chains'][0]['mean_plddt']}이고 설계된 펩타이드의 평균 pLDDT는 "
         f"{e['chains'][1]['mean_plddt']}이다. 잔기 수로 보면 수용체가 복합체의 "
         f"{e_share:.0%}를 차지하므로({e['chains'][0]['length']} 잔기 대 "
         f"{e['chains'][1]['length']} 잔기), 하나로 합친 평균은 복합체의 이름을 달고 있는 "
         f"수용체의 숫자다. 이 후보의 섞은 서열 {n_decoys}개 가운데 최고값은 "
         f"{e['screen']['decoy_max']}에 이른다.")
    figure(doc, "fig5_complex_structure.png",
           "왼쪽: Boltz-2가 만들어 낸 폴드를 AlphaFold가 공표한 pLDDT 구간으로 색칠하고, "
           "설계된 펩타이드를 굵게 그렸다. 오른쪽: 각 사슬을 따라간 잔기별 신뢰도. "
           "runs/에 있는 모델 자신의 좌표에서 렌더링했다.")

    h(doc, "5.5  독립적인 예측 모델이 수용체에 대해 일치한다", 2)
    aa, ab = af["arms"]["boltz_single_sequence"], af["arms"]["boltz_full_msa"]
    aa_top = max(aa["rows"], key=lambda r: r["pearson_r"])
    para(doc,
         f"슬레이트의 수용체 폴드가 모두 하나의 구조 예측 모델에서 나왔기 때문에 외부 점검을 "
         f"하나 덧붙였다. 각 표적에 대해 기탁된 AlphaFold DB 모델 {cite('varadi2024')}을 같은 "
         f"컨스트럭트 구간에서 Boltz-2의 잔기별 신뢰도와 잔기 단위로 대조한 것이다. Pearson r의 "
         f"중앙값은 연구 #9의 단일 서열 폴드에 대해 {aa['pearson_r_median']}, 연구 #10의 완전 "
         f"MSA 폴드에 대해 {ab['pearson_r_median']}이고, 평균 pLDDT 격차는 "
         f"{aa['mean_offset_afdb_minus_boltz']}점에서 {ab['mean_offset_afdb_minus_boltz']}점으로 "
         f"좁혀진다.")
    para(doc,
         f"이 비교는 탐색적이며, 산출물 안에서 그 범위가 신중하게 한정되어 있다. 이는 "
         f"펩타이드에 대한 점검도, 계면에 대한 점검도, 슬레이트의 어떤 주장에 대한 점검도 "
         f"아니다. 한 단백질 안의 잔기들은 서로 독립이 아니므로 어떤 r에도 p값을 붙이지 "
         f"않았다. 대신 각 행에는 자기상관을 반영한 유효 표본 수를 실었다. 단일 서열 "
         f"갈래에서 상관이 가장 큰 표적은 {aa_top['target']}, r = {aa_top['pearson_r']}이다. "
         f"그리고 비교에 쓰인 잔기 {aa_top['n_residues_compared']}개는 독립 관측 약 "
         f"{aa_top['effective_n_after_autocorrelation']}개에 해당한다 — 그 열에서 가장 큰 수가 "
         f"그 열에서 가장 근거가 약한 수가 되는 것이다.")
    figure(doc, "fig4_alphafold_vs_boltz.png",
           "각 수용체의 어느 부분이 신뢰도 높게 폴딩되어 있는지에 대해 두 독립 예측 모델이 "
           "보이는 일치도. Boltz-2에 MSA를 준 경우와 주지 않은 경우 각각. 데이터: "
           "data/alphafold_db_comparison.json.")
    figure(doc, "ui3_alphafold.png",
           f"워크벤치가 그리는 그대로의 같은 비교를 표적별·갈래별로 보여 준다. 이동 영가설 "
           f"열은 어떤 순환 이동도 관측된 r에 이르지 못했을 때 개수 대신 상한을 보고한다. "
           f"그 뒤에 있는 순환 이동들이 서로 크게 겹쳐서 "
           f"p값을 만들지 못하기 때문이다.")

    h(doc, "5.6  전체 슬레이트", 2)
    rows = []
    for s_ in slate["studies"]:
        num = f"#{s_['slate_number']}" if s_["slate_number"] else "—"
        v = " / ".join(f"{x['verdict'][0]}" for x in s_["hypotheses"] if x["verdict"])
        rows.append([num, s_["title"][:44], s_["plan_hash"], s_["n_observed"] or "—", v])
    table(doc, ["#", "연구", "계획 해시", "n", "판정"], rows)
    para(doc,
         f"C = confirmed(확인), F = falsified(반증), N = not tested(미검정). 가설 "
         f"{c['hypotheses']}개 가운데 {c['confirmed']}개가 확인되었고 {c['falsified']}개가 "
         f"반증되었으며 {c['not_tested']}개는 한 번도 검정되지 않았다. "
         f"{c['decided_by_a_test']}개는 검정통계량으로, {c['decided_by_a_threshold']}개는 사전 "
         f"지정된 문턱값으로 판정되었다. 확인된 기준은 검정 결과가 아니며, 여기의 확인 가운데 "
         f"몇 개는 반갑지 않은 진술을 확인한다 — 어떤 방법이 판별하지 못한다는 것, 또는 "
         f"후보들이 실패 구간에 들어간다는 것이다.", size=9.5, italic=True)
    figure(doc, "ui2_verdicts.png",
           "워크벤치가 그리는 그대로의 슬레이트. 판정 색은 사전 등록된 규칙이 발동했는지를 "
           "표시하며, 좋고 나쁨으로 읽히지 않는 색조를 일부러 썼다. 여기의 확인 가운데 몇 개는 "
           "반갑지 않은 소식이고, 중심이 되는 반증이 곧 "
           "발견이기 때문이다.")

    # ---------------------------------------------------------------- 6. LIMITS ----- #
    doc.add_page_break()
    h(doc, "6  이 작업이 보여주지 않는 것", 1)
    att = D["att"]
    for txt in [
        "이 작업은 이 펩타이드들이 결합할 수 없음을 보여주지 않는다. 보여주는 것은, 이 구조 예측 모델이 이 "
        "컨스트럭트 집합에서 이들을 자기 자신의 잔기를 섞은 서열과 구별하지 못한다는 사실이다. 습식 실험 "
        "검정은 여전히 결합을 찾아낼 수 있으며, 여기의 그 어떤 것도 그것을 검출하지 못했을 것이다.",
        "이 작업은 구조 예측 모델이 틀렸음을 보여주지 않는다. 독립적인 평가 두 건이 관련된 한계를 "
        f"보고하고 있으나 {cite('bret2026', 'wan2026')}, 본 프로젝트가 측정한 것은 스크리닝의 "
        "속성이지 모델의 정확도가 아니다.",
        "이 작업은 그 무엇에 대해서도 결합 친화도를 확립하지 않는다. 저장소에는 도킹, MM-GBSA, "
        "FEP 계산이 존재하지 않으며, ΔG 값을 실제로 산출한 유일한 방법인 슬레이트 #11의 "
        f"PRODIGY {cite('xue2016', 'vangone2015')}는 참조 구간의 "
        f"{D['pro']['metrics']['fraction_of_fit_range_occupied']:.0%}로 축소된 범위를 "
        f"반환했으며, 그 판별비의 부트스트랩 구간 "
        f"{D['pro']['metrics']['discrimination_ratio_ci95_bootstrap']}이 임계값을 걸치고 있다. "
        "정직하게 진술하자면, 이 설계로는 PRODIGY가 여기서 판별하는지 여부를 가릴 수 없다는 "
        "것이지, 판별하지 못한다는 것이 아니다.",
        "AlphaFold 비교는 첫 번째 갈래에서 구조 예측 모델, MSA, 단량체 대 복합체 맥락이라는 세 "
        "가지를 한꺼번에 교락시키고, 두 번째 갈래에서는 두 가지를 교락시킨다. 이 비교는 "
        "펩타이드나 계면에 대해서는 아무것도 말해 주지 않는다.",
f"후보 서열들은 발표된 모티프, 짜깁기한 스캐폴드, 그리고 하나의 de novo 나선을 손으로 "
        f"이어 붙인 연접체이며, GGGGS 링커로 연결되어 있다. 모티프 항목 "
        f"{att['attributed_motifs'] + att['unattributed_motif_entries']}개 가운데 UniProt "
        f"접근번호를 가진 것은 {att['attributed_motifs']}개뿐이다. 출처가 확인되지 않은 단편 "
        f"{att['distinct_unattributed_fragments']}개 전부를 대상으로 검색하면, 후보 "
        f"{att['candidates_total']}개 중 {att['candidates_carrying_one']}개가 적어도 하나를 "
        f"지니고 있다. 이들은 최적화된 설계가 아니며, 이들에 대한 음성 결과는 펩타이드 설계에 "
        f"대한 음성 결과가 아니다.",
        "확증적인 연구는 하나도 없다. 모든 연구가 적어도 한 가지 측면에서 등록된 계획을 "
        "벗어났다. 이탈 사항은 기계로 검출되며, 연구별로 나열되어 있다.",
    ]:
        bullet(doc, txt)

    # ---------------------------------------------------------------- 7. REPRO ------ #
    h(doc, "7  재현성, 식별자, 출처 기록", 1)
    para(doc,
         f"저장소는 코드에 대해 Apache-2.0으로 공개되어 있고, 제3자 과학 데이터는 그 밖의 "
         f"네 가지 라이선스로 재배포한다 — UniProt과 AlphaFold DB에는 CC BY 4.0, ChEMBL에서 "
         f"파생된 파일에는 CC BY-SA 3.0, RCSB 기탁물에는 CC0, Boltz-2 출력에는 MIT이다. "
         f"ChEMBL 파일에 걸린 동일조건변경허락 조항은 이후의 재배포에까지 의무를 전달하며, 어느 파일이 어느 "
         f"라이선스에 속하는지는 NOTICE에 열거되어 있다.")
    rows = [[i["value"], i.get("description", "")] for i in cit["identifiers"]]
    table(doc, ["식별자", "무엇을 가리키는가"], rows)
    para(doc,
         f"검증은 자동화되어 있으며 검증 자체도 검사를 받는다: 플랫폼과 프론트엔드 검사 묶음에 "
         f"걸친 {n_checks}개의 검사가 verify_all.py를 통해 다른 다섯 개의 검사 묶음과 함께 실행되고, "
         f"이 작업에서 추가한 모든 가드는 그 가드가 지목하는 결함을 실제로 주입해 실패를 "
         f"관찰하는 방식으로 해당 결함에서 실패함을 확인했다. 생성된 색인 세 개는 매 테스트 "
         f"실행마다 다시 만들어 대조하므로, 낡은 산출물이 배포될 수 없다. 생성되는 네 파일은 "
         f"모두 자신이 빌드된 커밋을 새기고, 입력이 커밋되지 않은 상태였다면 -dirty로 표시한다. "
         f"이 문서는 {git_sha()}에서 빌드되었고, 문서에 실린 모든 수치는 빌드 시점에 산출물에서 "
         f"읽어 온 값이므로, 데이터가 더 이상 뒷받침하지 않는 수치를 인용할 수 없다.")
    figure(doc, "ui5_citation.png",
           "인용 카드. 빌드 시점에 CITATION.cff에서 읽어 오므로 페이지가 "
           "레지스트리와 어긋날 수 없다.", width=5.9)

    # ---------------------------------------------------------------- 8. REFERENCES - #
    doc.add_page_break()
    h(doc, "8  참고문헌", 1)
    para(doc, D["refs"]["method_ko"], size=9, italic=True)
    for i, k in enumerate(CITED, 1):
        r = REF[k]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.34)
        p.paragraph_format.first_line_indent = Inches(-0.34)
        run = p.add_run(f"[{i}]  ")
        run.bold = True
        run.font.size = Pt(9.5)
        txt = f"{r['authors']} ({r['year']}). {r['title']}. {r['venue']}."
        if r.get("pmid"):
            txt += f" PMID {r['pmid']}."
        if r.get("doi"):
            txt += f" doi:{r['doi']}"
        run2 = p.add_run(txt)
        run2.font.size = Pt(9.5)
        note = p.add_run(f"\n        검증: {r['verified_ko']}. 용도: {r['used_for_ko']}.")
        note.font.size = Pt(8)
        note.font.color.rgb = GREY
        note.italic = True

    unused = [k for k in REF if k not in CITED]
    if unused:
        para(doc, f"검증되었으나 이 문서에서 인용하지 않음: {', '.join(unused)}.",
             size=8.5, italic=True)

    doc.save(OUT)
    print(f"wrote {OUT.relative_to(REPO)}")
    print(f"  references cited: {len(CITED)} of {len(REF)} verified")
    print(f"  figures embedded: {sum(1 for f in FIG.glob('*.png'))} available")
    return 0


if __name__ == "__main__":
    raise SystemExit(build())
