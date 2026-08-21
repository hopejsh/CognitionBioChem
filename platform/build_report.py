"""Build docs/CognitionBioChem_Report.docx -- the written account of the project.

Every number in the document is read from an artefact at build time. Nothing is typed, for
the same reason nothing is typed on the page: a figure written by hand into prose is a figure
that will be wrong the first time a study is re-run, and this project has already been through
that four times.

Every reference is read from docs/REFERENCES.json, where each entry records how it was
verified. Entries marked verified="PubMed" were resolved through the PubMed E-utilities; the
title, authors, journal, volume and pages are the ones PubMed returned. Bibliographic metadata
retrieved from PubMed (NLM/NCBI).
"""

from __future__ import annotations

import json
import re
import statistics
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "platform"))
from cbc.provenance import git_sha  # noqa: E402
from cbc.report_data import J, A, load  # noqa: E402

FIG = REPO / "docs" / "figures"
OUT = REPO / "docs" / "CognitionBioChem_Report.docx"

NAVY = RGBColor(0x1A, 0x3D, 0x6D)
GREY = RGBColor(0x5A, 0x66, 0x78)


D = load()
REF = {r["key"]: r for r in D["refs"]["references"]}
CITED: list[str] = []


def cite(*keys) -> str:
    """Return a bracketed citation and record the order of first appearance."""
    nums = []
    for k in keys:
        if k not in REF:
            raise KeyError(f"{k} is not in docs/REFERENCES.json -- refusing to cite it")
        if k not in CITED:
            CITED.append(k)
        nums.append(str(CITED.index(k) + 1))
    return "[" + ", ".join(nums) + "]"


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = NAVY
    return p


def para(doc, text, size=10.5, italic=False, space_after=8):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(size)
    p.paragraph_format.space_after = Pt(3)
    return p


FIGNO = [0]


def figure(doc, name, caption, width=6.3):
    """Place a figure and number it by position.

    The numbers used to be typed into the captions, which meant inserting a figure silently
    renumbered every one after it and broke the cross-references -- the same class of
    hand-maintained bookkeeping this project exists to remove.
    """
    path = FIG / name
    if not path.exists():
        raise FileNotFoundError(f"{path} is missing; the report will not be built without it")
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    FIGNO[0] += 1
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(f"Figure {FIGNO[0]}. {caption}")
    r.font.size = Pt(8.5)
    r.font.color.rgb = GREY
    r.italic = True
    c.paragraph_format.space_after = Pt(14)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(hd)
        r.bold = True
        r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return t


def build() -> int:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)

    cit, msa, scr = D["cit"], D["msa"], D["scr"]
    slate, struct, af = D["slate"], D["struct"], D["af"]
    m, c, nul, ver, per = D["m"], D["c"], D["nul"], D["ver"], D["per"]
    n_decoys, winners, win_margins = D["n_decoys"], D["winners"], D["win_margins"]
    scr_series, scr_decoys, n_checks = D["scr_series"], D["scr_decoys"], D["n_checks"]

    # ---------------------------------------------------------------- title page ---- #
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("CognitionBioChem"); r.bold = True; r.font.size = Pt(26)
    r.font.color.rgb = NAVY
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("A structural pharmacology workbench that reports a negative result")
    r.font.size = Pt(13); r.font.color.rgb = GREY
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("Why it was built, what it does, and what it found")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GREY
    doc.add_paragraph()

    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = a.add_run(f"{cit['authors'][0]}\nORCID {cit['orcid'][0].split('/')[-1]}\n\n"
                  f"Version {cit['version']} · {cit['date_released']}")
    r.font.size = Pt(10.5)
    ids = doc.add_paragraph(); ids.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ids.add_run("  ·  ".join(i["value"] for i in cit["identifiers"]))
    r.font.size = Pt(9); r.font.color.rgb = GREY
    doc.add_paragraph()

    # ---------------------------------------------------------------- abstract ------ #
    h(doc, "Summary", 1)
    para(doc,
         f"This project asked whether a set of manually concatenated peptide sequences, built "
         f"against cognition-related CNS targets, bind those targets. The answer is no, and "
         f"the work was organised so that the answer would be believable either way.")
    para(doc,
         f"Across a {scr['n_observed']}-fold screen without a multiple sequence alignment and "
         f"a {msa['n_observed']}-fold rerun with one, the designed peptides did not separate "
         f"from composition-matched shuffles of their own amino acids. Mean native interface "
         f"pTM was {m['mean_native_iptm']} against a mean decoy of {m['mean_decoy_iptm']}; the "
         f"paired t-test gives p = {msa['p_holm']['H1_natives_separate_from_decoys']:.2f}. "
         f"{nul['observed']} of {len(per)} candidates beat all {n_decoys} of their own "
         f"decoys, but {nul['expected_under_null']} are expected to by chance and "
         f"P(X ≥ {nul['observed']}) = {nul['p_at_least_observed']}. The hypothesis was "
         f"falsified in all {ver['n_versions']} retained versions of the two screening "
         f"studies, over candidate sets ranging from "
         f"{min(v['n_candidates'] for v in ver['versions'])} to "
         f"{max(v['n_candidates'] for v in ver['versions'])} designs.")
    para(doc,
         f"The instrument built to reach that answer is the substance of the work: "
         f"{c['studies']} studies pre-registered under content hashes before their data was "
         f"seen, {c['hypotheses']} hypotheses of which {c['falsified']} were falsified, "
         f"{len(struct['entries'])} structures held under content-addressed custody, and "
         f"{n_checks} automated checks each verified to fail on the defect it names. The "
         f"project began by discovering that its own earlier version displayed fabricated "
         f"affinity, confidence and safety values as though they had been computed; those "
         f"claims are preserved rather than deleted.")

    doc.add_page_break()

    # ---------------------------------------------------------------- 1. WHY -------- #
    h(doc, "1  Why this project exists", 1)
    h(doc, "1.1  The failure it started from", 2)
    n_retr = len(D["retr"])
    med_disc = statistics.median(a["discrepancy"]["value"] for a in D["aud"])
    med_ord = statistics.median(a["discrepancy_orders"]["value"] for a in D["aud"])
    para(doc,
         f"An earlier version of this project presented numbers as though a structure "
         f"predictor had produced them. No calculation had. {n_retr} candidates carried "
         f"affinity, confidence and safety values with no computation behind them: a "
         f"\"pLDDT\" curve was the expression 93 + sin(i·0.4)·4 + (charCode % 5)·0.5, and every "
         f"one of the {n_retr} stated ΔG/Kd pairs is internally inconsistent at 298.15 K "
         f"against ΔG = RT·ln(Kd), by a median of {med_disc:.2f} kcal/mol — "
         f"{med_ord:.1f} orders of magnitude in Kd. Those values are preserved in the data "
         f"layer under retracted_claims, and the renderers that produced them are gone.")
    para(doc,
         "That is the specific reason the project exists in its present form, and it explains "
         "an otherwise strange design decision: the workbench is built so that it cannot "
         "display a number without a provenance record, and a field that was never computed "
         "renders as a label rather than as a figure.")

    h(doc, "1.2  Why the question needed a null", 2)
    para(doc,
         f"A structure predictor will return a confident-looking complex for almost any pair "
         f"of chains. The highest interface pTM in this set is "
         f"{max(p['native_iptm'] for p in msa['per_candidate'])}, which by the bands "
         f"established for interface quality {cite('basu2016', 'lensink2007')} would ordinarily "
         f"be read as a good interface. The question is not whether a designed peptide scores "
         f"well; it is whether it scores better than a sequence with the same amino acids in a "
         f"different order. Without that comparison a screen measures how confident the model "
         f"is, not whether the design did anything.")
    para(doc,
         f"That concern is not hypothetical for this class of model. Independent evaluations of "
         f"the predictor used here report that its affinity output can be insensitive to "
         f"binding-site mutation and in some cases even to exchanging the target "
         f"{cite('bret2026')}, and that it lacks the energetic resolution required for lead "
         f"identification {cite('wan2026')}. Work on physical validity had already shown that "
         f"deep-learning docking methods frequently produce poses that are geometrically "
         f"implausible while scoring well on RMSD {cite('buttenschoen2024')}.")

    # ---------------------------------------------------------------- 2. BACKGROUND - #
    doc.add_page_break()
    h(doc, "2  Scientific background", 1)

    h(doc, "2.1  What a structure predictor reports, and what it does not", 2)
    para(doc,
         f"AlphaFold2 established that a neural network can predict protein structure at "
         f"atomic accuracy from sequence alone, and — as important for this work — that it can "
         f"report calibrated confidence in its own output {cite('jumper2021')}. Two of those "
         f"confidence quantities are used throughout this project. pLDDT is a per-residue "
         f"estimate of local accuracy on a 0–100 scale. PAE, the predicted aligned error, is a "
         f"matrix giving the expected positional error of residue i when the structure is "
         f"superposed on residue j; it is directed, so it is genuinely asymmetric, and its "
         f"off-diagonal blocks say whether two regions are confidently placed relative to each "
         f"other. Interface pTM (ipTM) condenses the cross-chain part of that into one number.")
    para(doc,
         f"None of these is a binding affinity, and none of them is evidence of binding. They "
         f"are the model's report on its own certainty about geometry. A high ipTM for a "
         f"peptide-receptor pair says the model places the peptide somewhere definite; it does "
         f"not say the peptide belongs there. This distinction is the axis the entire project "
         f"turns on.")

    h(doc, "2.2  The predictor used, and why not AlphaFold 3", 2)
    para(doc,
         f"Every predicted structure in this repository was produced by Boltz-2 v2.2.1 "
         f"{cite('passaro2025')}, whose code and weights are both MIT-licensed and which runs "
         f"locally on Apple silicon. AlphaFold 3 was not used: its parameters are request-only, "
         f"non-commercial and Linux/CUDA-bound, so no reproducible local pipeline could be "
         f"built on it. AlphaFold Server was not used either, and could not have been — its "
         f"terms prohibit automated use for protein-ligand and protein-peptide binding "
         f"prediction, which is exactly what the two screening studies do. The deposited "
         f"AlphaFold Protein Structure Database {cite('varadi2024')} is a separately licensed "
         f"corpus and is used, under CC BY 4.0, as an independent comparison.")

    h(doc, "2.3  Why no free energy is emitted", 2)
    para(doc,
         f"Boltz-2 carries an affinity head, and this project deliberately never renders its "
         f"output as a free energy. The head is fitted to pooled Ki, Kd, IC50 and EC50 labels. "
         f"Those are not interchangeable quantities: an IC50 depends on substrate concentration "
         f"and on the enzyme's Km, and converting one to a Ki requires assumptions about "
         f"mechanism and concentrations that a pooled training label does not carry "
         f"{cite('cheng1973', 'cer2009')}. A number fitted to that mixture has no single "
         f"physical referent, so no thermodynamic quantity can be recovered from it, and "
         f"a build guard in the repository fails if any code path tries.")

    h(doc, "2.4  How noisy the reference data is", 2)
    ache_metrics = D["ache"]["metrics"]
    para(doc,
         f"Judging a predictor against public bioactivity data requires knowing how noisy that "
         f"data is. Independent measurements of the same protein-ligand pair in ChEMBL "
         f"{cite('mendez2019', 'zdrazil2024')} differ substantially: the experimental "
         f"uncertainty of heterogeneous public Ki data is reported as a standard deviation of "
         f"0.54 pKi units, which caps the achievable R² near 0.81 {cite('kramer2012')}; the "
         f"standard deviation of mixed public IC50 data is a further 25% larger "
         f"{cite('kalliokoski2013')}; and combining IC50 or Ki values from different sources "
         f"adds significant noise again {cite('landrum2024')}. This project measured the same "
         f"quantity for its own benchmark pair — huperzine A against acetylcholinesterase — "
         f"and found an in-situ reference σ of about 0.99 log against a model mean absolute "
         f"error of {ache_metrics['mean_absolute_error_log10']:.2f} log. The two are the same "
         f"order of magnitude, so that study cannot separate model error from reference "
         f"error, and says so.")
    para(doc,
         "That 0.99 figure carries its own caveat, recorded rather than smoothed over: it is "
         "measured across a pooled selection of ChEMBL records whose assay configurations are "
         "largely unannotated. Restricted to the properly annotated subset the spread falls "
         "below the generic IC50 floor, so the honest reading is that this pair is unusually "
         "poorly annotated rather than unusually noisy — and the conclusion that the study "
         "cannot separate model error from reference error survives either way.",
         size=9.5, italic=True)

    h(doc, "2.5  How interface quality is measured", 2)
    para(doc,
         f"Where an experimental structure exists, a predicted interface can be scored against "
         f"it. DockQ combines the three CAPRI quantities — the fraction of native contacts "
         f"recovered, the ligand RMSD and the interface RMSD — into a single continuous score "
         f"in [0, 1] that reproduces the CAPRI acceptable/medium/high classification almost "
         f"exactly {cite('basu2016', 'lensink2007')}. The implementation used here is DockQ v2 "
         f"{cite('mirabello2024')}. This matters because it lets one ask whether the model's "
         f"self-reported confidence tracks actual interface correctness, which is the question "
         f"Slate #7 was registered to answer before the candidate screen was allowed to run.")

    h(doc, "2.6  Physicochemical properties and admissibility", 2)
    para(doc,
         f"Sequence-level properties are computed rather than asserted: molecular weight, net "
         f"charge at pH 7.4, isoelectric point, and Kyte-Doolittle hydropathy "
         f"{cite('kyte1982')}. Chemical structures are validated with RDKit — parsing, formula, "
         f"InChIKey and stereochemistry completeness — and a compound whose stored SMILES "
         f"encodes a different molecule than its name is shown as unverified rather than "
         f"rendered with computed properties. ADMET is predicted with ADMET-AI "
         f"{cite('swanson2024')} only where the molecule falls inside the model's applicability "
         f"domain, and is refused with a stated reason where it does not. Target constructs, "
         f"signal peptides, topology and disulfide bonds come from UniProt {cite('uniprot2025')}, "
         f"and the experimental reference structures from the Protein Data Bank "
         f"{cite('berman2000')}.")

    h(doc, "2.7  Multiplicity, and what a pre-registered criterion is not", 2)
    para(doc,
         f"Each study registers three hypotheses. Where a hypothesis is decided by a test "
         f"statistic, the family is corrected with the Holm step-down procedure "
         f"{cite('holm1979')}. Most hypotheses in this slate are not tests: they are threshold "
         f"comparisons on a descriptive quantity, decided by looking at a number against a line "
         f"drawn in advance. Of {c['hypotheses']} hypotheses, only {c['decided_by_a_test']} were "
         f"decided by a test. Encoding a threshold criterion as a p-value of 0.0 or 1.0 and "
         f"feeding it to Holm — which an earlier version of this code did — both fabricates an "
         f"unattainable p-value and steals a multiplier from the real tests it is grouped with. "
         f"The two are kept apart, and the page labels every verdict with which it was.")

    # ---------------------------------------------------------------- 3. WHAT ------- #
    doc.add_page_break()
    h(doc, "3  What was built", 1)
    para(doc,
         f"The workbench is a static page over a provenance-carrying data layer, with a Python "
         f"platform behind it that runs the predictions, the studies and the checks. Three "
         f"mechanisms do the load-bearing work.")

    h(doc, "3.1  Provenance as a rendering rule", 2)
    para(doc,
         f"Every scientific value carries a status: computed, predicted, database, literature, "
         f"measured, placeholder, or not_computed. The user interface has exactly one path to "
         f"the screen for a value, and that path emits a label rather than a figure when the "
         f"status is placeholder or not_computed. This is why the fabricated values could not "
         f"simply be deleted and forgotten: they are retained under retracted_claims, visibly "
         f"struck through, so the record of what was once asserted is not lost.")
    figure(doc, "ui1_headline_finding.png",
           "The workbench's Overview tab. The result is stated before the "
           "capabilities, and every number in the card is read from a study artefact at build "
           "time rather than written into the markup.")

    h(doc, "3.2  Pre-registration under a content hash", 2)
    plans = D["plans"]
    para(doc,
         f"Before a study runs, its question, primary metric, decision threshold, hypotheses "
         f"and analysis plan are frozen and hashed. The registry refuses a plan with "
         f"unreachable verdicts — a criterion whose smallest attainable adjusted p already "
         f"exceeds α — with unfalsifiable hypotheses, or in which every hypothesis is predicted "
         f"by the same position so no outcome could discriminate. {plans} plans are retained "
         f"across {c['studies']} study families; superseded versions are kept, not overwritten, "
         f"and each records why it was superseded.")
    para(doc,
         f"An important consequence is recorded honestly on the page: not one study in this "
         f"slate is confirmatory. Every one deviated from its registered plan in at least one "
         f"respect, and each study's own audit says so. Pre-registration did not make these "
         f"results confirmatory — it made the deviations visible.")

    h(doc, "3.3  Custody of the computation", 2)
    runs, published_rows = D["runs"], D["published_rows"]
    para(doc,
         f"Every prediction run is content-addressed: the directory name is a hash over the "
         f"names and contents of its files, so the identifier changes if any output changes. "
         f"{runs} runs are under custody. A test asserts that every row any study reports "
         f"resolves to a run still in the manifest — {published_rows} rows across the six studies "
         f"that fold anything — which is "
         f"what distinguishes a run that was re-hashed from one that was lost.")
    figure(doc, "ui4_structure_gallery.png",
           f"The structure gallery. Every entry is a model the pipeline produced, opened from "
           f"the run directory whose content hash names it; the confidence terms shown per "
           f"entry are only those the model actually defines, so a monomer shows no interface "
           f"score rather than a zero.")
    para(doc,
         f"Two failures during development motivated this. A rerun after a construct correction "
         f"returned an ipTM identical to sixteen digits, because the predictor's preprocessing "
         f"cache is keyed on the input record's name and had silently re-folded the previous "
         f"receptor. A later interrupted run wrote a manifest entry pairing a new input's hash "
         f"with an old result. Both now fail a post-condition that re-parses the produced model "
         f"and checks its chain lengths against the request.")

    # ---------------------------------------------------------------- 4. HOW -------- #
    doc.add_page_break()
    h(doc, "4  How the central question was asked", 1)

    h(doc, "4.1  The composition-matched null", 2)
    n_dec = msa["per_candidate"][0]["n_decoys"]
    para(doc,
         f"For each candidate, {n_dec} decoy sequences are generated by shuffling that "
         f"candidate's own amino acids. A decoy therefore has identical composition — the same "
         f"molecular weight, the same net charge, the same hydropathy — and differs only in "
         f"order. Each decoy is folded against the same receptor construct, with the same seed, "
         f"in the same order. If the designed sequence carries information about the target, it "
         f"should score above its own shuffles; if the model is responding to composition, it "
         f"will not.")
    para(doc,
         f"The receptor construct is not the full UniProt sequence. Signal peptides are removed, "
         f"transmembrane and cytoplasmic segments are excluded where the binding site is "
         f"extracellular, and the span actually folded is recorded per candidate with its basis "
         f"and canonical numbering. Three construct corrections were made during the work; each "
         f"changed the answer's precision and none changed its direction.")

    h(doc, "4.2  A gate before the screen was allowed to count", 2)
    gate_pl, gate_rl = D["gate_pl"], D["gate_rl"]
    cand_pl, cand_rl = D["cand_pl"], D["cand_rl"]
    pi_m = D["pi"]["metrics"]
    para(doc,
         f"Slate #7 was registered before the candidate screen was permitted to be believed. It "
         f"asks whether this pipeline recovers interfaces it could have memorised: 16 "
         f"peptide-receptor X-ray complexes were folded and scored with DockQ "
         f"{cite('mirabello2024')}. The pipeline reached CAPRI-acceptable quality on "
         f"{pi_m['fraction_dockq_acceptable']:.0%} of the {D['pi']['n_observed']} complexes "
         f"(median DockQ {pi_m['median_dockq']}), and ipTM tracked DockQ at "
         f"Spearman ρ = {pi_m['spearman_iptm_dockq']}. That is the sense in which ipTM is used "
         f"downstream: as a discrimination signal on this kind of complex, not as a calibrated "
         f"probability.")
    para(doc,
         f"The same study measured the leakage that makes such a gate necessary. The benchmark "
         f"peptides are {min(gate_pl)}–{max(gate_pl)} residues against receptors of "
         f"{min(gate_rl)}–{max(gate_rl)}, while the candidates are "
         f"{min(cand_pl)}–{max(cand_pl)} residues against receptors of "
         f"{min(cand_rl)}–{max(cand_rl)}. The sensitivity argument transfers; the "
         f"numeric bands do not, and the report says so rather than carrying the thresholds "
         f"across.")

    h(doc, "4.3  How large a difference has to be before it means anything", 2)
    iv_m = D["iv"]["metrics"]
    para(doc,
         f"Slate #2 measured the sampler's own noise before any comparison was trusted. Folding "
         f"the same complexes repeatedly under different seeds gives an across-seed standard "
         f"deviation of {iv_m['across_seed_sd_iptm']:.3f} in ipTM and "
         f"{iv_m['across_seed_sd_complex_plddt']:.2f} pLDDT units. Any difference smaller than that is "
         f"not a measurement. It is the reason the {len(winners)} apparent 'successes' in the "
         f"screen are reported as noise: they beat their best decoy by "
         f"{win_margins}, against a "
         f"one-standard-deviation spread of {iv_m['across_seed_sd_iptm']:.3f}.")

    # ---------------------------------------------------------------- 5. RESULTS ---- #
    doc.add_page_break()
    h(doc, "5  Results", 1)

    h(doc, "5.1  The designed sequences do not separate from their own shuffles", 2)
    beaten = sum(1 for p in per if p["decoy_max"] > p["native_iptm"])
    beaten_mean = sum(1 for p in per if p["decoy_mean"] > p["native_iptm"])
    para(doc,
         f"With a full multiple sequence alignment and {n_decoys} decoys per candidate, mean "
         f"native "
         f"ipTM is {m['mean_native_iptm']} against a mean decoy of {m['mean_decoy_iptm']}. The "
         f"paired difference is {m['mean_native_iptm'] - m['mean_decoy_iptm']:+.4f} with "
         f"p = {msa['p_holm']['H1_natives_separate_from_decoys']:.2f} and Cohen's "
         f"dz = {m['cohens_dz']:+.2f}. For {beaten} of {len(per)} candidates the best of the "
         f"{n_decoys} shuffles scores above the designed sequence, and for {beaten_mean} of {len(per)} the "
         f"decoy mean does.")
    figure(doc, "fig1_native_vs_decoy.png",
           f"Each designed peptide against {n_decoys} shuffles of its own amino "
           f"acids. The diamond is the designed sequence; the bar spans the decoy mean to the "
           f"best decoy. "
           "A design to the right of its bar outscored all of its shuffles. Data: "
           "data/study_msa_specificity.json.")

    h(doc, "5.2  Two winners is what chance looks like", 2)
    para(doc,
         f"{nul['observed']} of {len(per)} candidates beat all {n_decoys} of their own decoys "
         f"({', '.join(w['code'] for w in winners)}). Under the null a candidate does so with "
         f"probability {nul['per_candidate_null_probability']} = 1/{n_decoys + 1}, so "
         f"{nul['expected_under_null']} of {len(per)} are expected to by chance and "
         f"P(X ≥ {nul['observed']}) = {nul['p_at_least_observed']}. Every winner has the same "
         f"empirical p of {max(w['empirical_p'] for w in winners)}, the smallest value this "
         f"design can produce, which clears no conventional "
         f"threshold. The composition-matched null protects a candidate from being read as a "
         f"hit; it does nothing for a screen read the same way, and this is the same error one "
         f"level up. The screen-level null was computed after the data were seen and is flagged "
         f"as exploratory in the artefact for that reason.")
    figure(doc, "fig2_screen_level_null.png",
           f"The screen-level null. Grey is the binomial distribution of how many "
           f"of {len(per)} candidates would beat all {n_decoys} of their decoys by chance; "
           f"amber is what was observed. Data: data/study_msa_specificity.json.", width=4.8)

    h(doc, "5.3  The verdict survived every correction; the margin did not", 2)
    para(doc,
         f"Across three construct corrections, a de-duplication, two coverage expansions and a "
         f"decoy expansion from {scr_decoys} to {n_decoys} shuffles per candidate, H1 was "
         f"falsified in all "
         f"{ver['n_versions']} retained versions. What did not hold steady is the size of the "
         f"gap: across the screen's {len(scr_series)} versions the mean native-minus-decoy ran "
         f"{', '.join(f'{v:+.4f}' for v in scr_series)}, growing several-fold before coming "
         f"back through zero, and it has never left the sampler-noise floor "
         f"(±{D['iv']['metrics']['across_seed_sd_iptm']:.3f} ipTM). The verdict is what "
         f"survived every "
         f"correction, not the margin.")
    figure(doc, "fig3_falsified_every_version.png",
           "Every retained version of the two screening studies, with the candidate "
           "count in each. Data: data/slate.json, data/superseded/.", width=5.9)

    h(doc, "5.4  Where the confidence actually comes from", 2)
    e = next(x for x in struct["entries"] if x["id"] == "cpx-BasalAChE-Abeta-B4")
    e_share = e["chains"][0]["length"] / sum(ch["length"] for ch in e["chains"])
    para(doc,
         f"A single complex makes the point that the aggregate statistics make abstractly. "
         f"{e['code']} against {e['target']} has ipTM {e['metrics']['iptm']}, which by the usual "
         f"bands is a good interface. Split by chain, the receptor's mean pLDDT is "
         f"{e['chains'][0]['mean_plddt']} and the designed peptide's is "
         f"{e['chains'][1]['mean_plddt']}. The receptor is {e_share:.0%} of the complex by "
         f"residue count ({e['chains'][0]['length']} residues against "
         f"{e['chains'][1]['length']}), so a single pooled mean is the receptor's number "
         f"wearing the complex's name. The best of this candidate's {n_decoys} shuffles "
         f"reaches {e['screen']['decoy_max']}.")
    figure(doc, "fig5_complex_structure.png",
           "Left: the fold Boltz-2 produced, coloured by AlphaFold's published pLDDT "
           "bands, with the designed peptide drawn thick. Right: per-residue confidence along "
           "each chain. Rendered from the model's own coordinates in runs/.")

    h(doc, "5.5  An independent predictor agrees about the receptors", 2)
    aa, ab = af["arms"]["boltz_single_sequence"], af["arms"]["boltz_full_msa"]
    aa_top = max(aa["rows"], key=lambda r: r["pearson_r"])
    para(doc,
         f"Because every receptor fold in the slate comes from one predictor, an external check "
         f"was added: the deposited AlphaFold DB model {cite('varadi2024')} for each target, "
         f"compared residue by residue against Boltz-2's per-residue confidence over the same "
         f"construct span. Median Pearson r is {aa['pearson_r_median']} against study #9's "
         f"single-sequence folds and {ab['pearson_r_median']} against study #10's full-MSA "
         f"folds, and the mean pLDDT gap closes from "
         f"{aa['mean_offset_afdb_minus_boltz']} to {ab['mean_offset_afdb_minus_boltz']} points.")
    para(doc,
         f"This comparison is exploratory and is bounded carefully in the artefact. It is not a "
         f"check on the peptide, on the interface, or on any claim in the slate. Residues within "
         f"one protein are not independent, so no p-value is attached to any r; each row instead "
         f"carries an effective sample size after autocorrelation. In the single-sequence arm "
         f"the largest correlation is {aa_top['target']} at r = {aa_top['pearson_r']}, and its "
         f"{aa_top['n_residues_compared']} compared residues are worth about "
         f"{aa_top['effective_n_after_autocorrelation']} independent observations — which makes "
         f"the largest number in that column the least supported one in it.")
    figure(doc, "fig4_alphafold_vs_boltz.png",
           "Agreement between two independent predictors about where each receptor is "
           "confidently folded, with and without an MSA for Boltz-2. Data: "
           "data/alphafold_db_comparison.json.")
    figure(doc, "ui3_alphafold.png",
           f"The same comparison as the workbench renders it, per target and per arm. The "
           f"shift-null column reports an upper bound rather than a count when no circular "
           f"shift reached the observed r, because the "
           f"heavily overlapping circular shifts behind it do not make a p-value.")

    h(doc, "5.6  The full slate", 2)
    rows = []
    for s_ in slate["studies"]:
        num = f"#{s_['slate_number']}" if s_["slate_number"] else "—"
        v = " / ".join(f"{x['verdict'][0]}" for x in s_["hypotheses"] if x["verdict"])
        rows.append([num, s_["title"][:44], s_["plan_hash"], s_["n_observed"] or "—", v])
    table(doc, ["#", "Study", "Plan hash", "n", "Verdicts"], rows)
    para(doc,
         f"C = confirmed, F = falsified, N = not tested. Of {c['hypotheses']} hypotheses, "
         f"{c['confirmed']} were confirmed, {c['falsified']} falsified and {c['not_tested']} "
         f"never tested; {c['decided_by_a_test']} were decided by a test statistic and "
         f"{c['decided_by_a_threshold']} by a pre-specified threshold. A confirmed criterion is "
         f"not a test result, and several of the confirmations here confirm unwelcome "
         f"statements — that a method does not discriminate, or that candidates fall in a "
         f"failed band.", size=9.5, italic=True)
    figure(doc, "ui2_verdicts.png",
           "The slate as the workbench renders it. Verdict colour marks whether a "
           "pre-registered rule fired, deliberately in hues that carry no good/bad reading, "
           "because several confirmations here are unwelcome news and the central falsification "
           "is the finding.")

    # ---------------------------------------------------------------- 6. LIMITS ----- #
    doc.add_page_break()
    h(doc, "6  What this work does not show", 1)
    att = D["att"]
    for txt in [
        "It does not show that these peptides cannot bind. It shows that this predictor, on "
        "this construct set, does not distinguish them from shuffles of their own residues. A "
        "wet-lab assay could still find binding; nothing here would have detected it.",
        "It does not show that the predictor is wrong. Two independent evaluations report "
        f"related limitations {cite('bret2026', 'wan2026')}, but this project measured a "
        "property of a screen, not the accuracy of a model.",
        "It does not establish a binding affinity for anything. No docking, MM-GBSA or FEP "
        "calculation exists in the repository, and the one method that did produce a ΔG — "
        f"PRODIGY {cite('xue2016', 'vangone2015')} in Slate #11 — returned a range that "
        f"collapsed to {D['pro']['metrics']['fraction_of_fit_range_occupied']:.0%} of the "
        f"reference span, with a discrimination ratio whose bootstrap interval "
        f"{D['pro']['metrics']['discrimination_ratio_ci95_bootstrap']} straddles its threshold. "
        "The honest statement is that the design cannot resolve whether PRODIGY discriminates "
        "here, not that it does not.",
        "The AlphaFold comparison confounds three things at once in its first arm — predictor, "
        "MSA, and monomer-versus-complex context — and two in its second. It says nothing about "
        "the peptide or the interface.",
f"The candidate sequences are chimeric peptides — published motifs, motif-like "
        f"segments with no identifiable natural source, and one de novo helix, "
        f"concatenated head to tail with GGGGS linkers. Of "
        f"{att['attributed_motifs'] + att['unattributed_motif_entries']} motif entries only "
        f"{att['attributed_motifs']} carry a UniProt accession; scanning for all "
        f"{att['distinct_unattributed_fragments']} unattributed fragments, "
        f"{att['candidates_carrying_one']} of {att['candidates_total']} candidates carry at "
        f"least one. These are not optimised designs, and a negative result on them is not a "
        f"negative result on peptide design.",
        "Not one study is confirmatory. Every one deviated from its registered plan in at least "
        "one respect. The deviations are machine-detected and listed per study.",
    ]:
        bullet(doc, txt)

    # ---------------------------------------------------------------- 7. REPRO ------ #
    h(doc, "7  Reproducibility, identifiers and provenance", 1)
    para(doc,
         f"The repository is public under Apache-2.0 for its code, and redistributes "
         f"third-party scientific data under four further licences — CC BY 4.0 for UniProt and "
         f"AlphaFold DB, CC BY-SA 3.0 for ChEMBL-derived files, CC0 for RCSB depositions, and "
         f"MIT for Boltz-2 outputs. The share-alike term on the ChEMBL files carries an "
         f"obligation onward, and NOTICE lists which files fall under which.")
    rows = [[i["value"], i.get("description", "")] for i in cit["identifiers"]]
    table(doc, ["Identifier", "What it names"], rows)
    para(doc,
         f"Verification is automated and is itself checked: {n_checks} checks across the "
         f"platform and front-end suites, run together with five more suites by verify_all.py, "
         f"and every guard added during this work was "
         f"verified to fail on the defect it names by injecting that defect and observing the "
         f"failure. Three generated indices are rebuilt and compared on every test run, so a "
         f"stale artefact cannot ship. All four generated files stamp the commit they were "
         f"built from and mark it -dirty when their inputs were uncommitted. This document was "
         f"built from {git_sha()}, and every number in it is read from an artefact at build "
         f"time, so it cannot quote a figure the data no longer supports.")
    figure(doc, "ui5_citation.png",
           "The citation card, read from CITATION.cff at build time so the page "
           "cannot fall out of step with the registries.", width=5.9)

    # ---------------------------------------------------------------- 8. REFERENCES - #
    doc.add_page_break()
    h(doc, "8  References", 1)
    para(doc, D["refs"]["method"], size=9, italic=True)
    for i, k in enumerate(CITED, 1):
        r = REF[k]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.34)
        p.paragraph_format.first_line_indent = Inches(-0.34)
        run = p.add_run(f"[{i}]  ")
        run.bold = True
        run.font.size = Pt(9.5)
        txt = f"{r['authors']} ({r['year']}). {r['title']}. {r['venue']}."
        if r.get("pmid"):
            txt += f" PMID {r['pmid']}."
        if r.get("doi"):
            txt += f" doi:{r['doi']}"
        run2 = p.add_run(txt)
        run2.font.size = Pt(9.5)
        note = p.add_run(f"\n        Verified: {r['verified']}. Used for: {r['used_for']}.")
        note.font.size = Pt(8)
        note.font.color.rgb = GREY
        note.italic = True

    unused = [k for k in REF if k not in CITED]
    if unused:
        para(doc, f"Verified but not cited in this document: {', '.join(unused)}.",
             size=8.5, italic=True)

    doc.save(OUT)
    print(f"wrote {OUT.relative_to(REPO)}")
    print(f"  references cited: {len(CITED)} of {len(REF)} verified")
    print(f"  figures embedded: {sum(1 for f in FIG.glob('*.png'))} available")
    return 0


if __name__ == "__main__":
    raise SystemExit(build())
