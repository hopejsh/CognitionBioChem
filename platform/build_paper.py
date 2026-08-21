#!/usr/bin/env python3
"""Build the manuscript as a Word document, in English or Korean.

    ./.venv/bin/python platform/build_paper.py          -> docs/CognitionBioChem_Paper.docx
    ./.venv/bin/python platform/build_paper.py --ko     -> docs/CognitionBioChem_Paper_KO.docx

One renderer serves both editions. The alternative -- a second module with the strings
swapped -- is a copy, and a copy drifts: the first time a figure is added to one and not the
other, two documents claim to be the same paper and are not.

The prose lives once, in the drafted sections under the paper working directory; cbc.paper
parses it into format-neutral blocks and resolves every [@key] against the verified reference
library. This module only decides how a block looks on a page. The Korean edition renders the
same block list, so the two cannot disagree about a citation number or a figure number.

Every reference was retrieved through the PubMed E-utilities by the agent that proposed it,
re-fetched by an independent checker, and re-fetched a third time directly from NCBI esummary
and compared on title and first author before being admitted. Bibliographic metadata retrieved
from PubMed (NLM/NCBI).
"""

from __future__ import annotations

import sys

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(REPO / "platform"))
from cbc.paper import ORDER, parse, plain, render_reference  # noqa: E402
from cbc.provenance import git_sha  # noqa: E402
from cbc.report_data import load  # noqa: E402

SRC = REPO / "paper"
FIG = REPO / "docs" / "figures"
NAVY = RGBColor(0x1A, 0x3D, 0x6D)
GREY = RGBColor(0x55, 0x5F, 0x6C)

#: Per-edition typography and fixed labels. Word resolves a font per script, not per run, so
#: the Korean edition stamps one family on all four script slots; left to the default, Hangul
#: silently falls back to whatever the system offers and the page ends up in two unrelated
#: typefaces.
LANG = {
    "en": {
        "out": "CognitionBioChem_Paper.docx",
        "body": "Cambria", "sans": "Calibri", "mono": "Menlo", "east_asia": None,
        "abstract": "Abstract", "keywords": "Keywords: ",
        "references": "References", "availability": "Data and code availability",
        "figure": "Figure {n}. ",
        "preprint": "  [preprint; not PubMed-indexed]",
        "stamp": "Manuscript generated from commit {sha} · {date}",
        "refs_note": (
            "Every entry below was retrieved through the PubMed E-utilities by the specialist "
            "that proposed it, re-fetched independently by a second checker, and re-fetched a "
            "third time directly from NCBI esummary and compared on title and first author "
            "before being admitted. Titles, authors, journals, volumes, pages, years and DOIs "
            "are NCBI's. Four entries are preprints that PubMed does not index and say so. "
            "Bibliographic metadata retrieved from PubMed (NLM/NCBI)."),
    },
    "ko": {
        "out": "CognitionBioChem_Paper_KO.docx",
        "body": "Apple SD Gothic Neo", "sans": "Apple SD Gothic Neo", "mono": "Menlo",
        "east_asia": "Apple SD Gothic Neo",
        "abstract": "초록", "keywords": "주제어: ",
        "references": "참고문헌", "availability": "데이터 및 코드 공개",
        "figure": "그림 {n}. ",
        "preprint": "  [프리프린트, PubMed 미색인]",
        "stamp": "커밋 {sha}에서 생성 · {date}",
        "refs_note": (
            "아래 모든 항목은 그것을 제안한 전문가가 PubMed E-utilities로 직접 조회했고, 독립적인 "
            "검증자가 다시 받아왔으며, 세 번째로 NCBI esummary에서 직접 다시 받아 제목과 제1저자를 "
            "대조한 뒤에야 채택되었다. 제목·저자·학술지·권·페이지·연도·DOI는 NCBI가 돌려준 값이다. "
            "네 항목은 PubMed에 색인되지 않은 프리프린트이며 그 사실을 함께 밝힌다. 참고문헌 자체는 "
            "번역하지 않는다 — 인용은 찾을 수 있어야 하기 때문이다. "
            "서지 정보는 PubMed(NLM/NCBI)에서 가져왔다."),
    },
}

CFG: dict = {}


def _stamp(run, name: str):
    """Set a run's face on every script slot, not just the Latin one."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for slot in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(slot), name)
    return run


def runs_into(par, runs, *, size=11.0, color=None, italic=False, bold=False):
    for r in runs:
        run = par.add_run(r["t"])
        f = run.font
        f.size = Pt(size)
        _stamp(run, CFG["mono"] if r.get("mono") else CFG["body"])
        f.bold = bold or bool(r.get("bold"))
        f.italic = italic or bool(r.get("italic"))
        if color is not None:
            f.color.rgb = color
        if r.get("cite"):
            f.color.rgb = NAVY
            f.superscript = True
    return par


def build(lang: str = "en") -> int:
    global CFG
    CFG = LANG[lang]
    OUT = REPO / "docs" / CFG["out"]
    suffix = "" if lang == "en" else f".{lang}"
    paper = parse(SRC, SRC / "REFERENCES_PAPER.json", lang=lang)
    D = load(run_suites=False)
    cit = D["cit"]

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = CFG["body"]
    st.font.size = Pt(11)
    if CFG["east_asia"]:
        st.element.rPr.rFonts.set(qn("w:eastAsia"), CFG["east_asia"])
    for s in doc.sections:
        s.left_margin = s.right_margin = Inches(1.0)
        s.top_margin = s.bottom_margin = Inches(0.9)

    def para(text="", *, size=11.0, italic=False, bold=False, color=None,
             align=WD_ALIGN_PARAGRAPH.LEFT, after=8, first_line=None):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        if first_line is not None:
            p.paragraph_format.first_line_indent = Inches(first_line)
        if text:
            r = p.add_run(text)
            r.font.size = Pt(size)
            _stamp(r, CFG["body"])
            r.italic = italic
            r.bold = bold
            if color is not None:
                r.font.color.rgb = color
        return p

    # ------------------------------------------------------------------ title page --- #
    title_file = SRC / f"sec_title{suffix}.md"
    if not title_file.exists():
        raise FileNotFoundError(f"{title_file} is missing; the paper has no title")
    lines = [x.strip() for x in title_file.read_text().strip().split("\n") if x.strip()]
    t = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    r = t.add_run(lines[0])
    r.bold = True
    r.font.size = Pt(17)
    _stamp(r, CFG["sans"])
    r.font.color.rgb = NAVY
    if len(lines) > 1:
        para(lines[1], size=11.5, italic=True, color=GREY,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=16)

    a = para(align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    r = a.add_run(cit["authors"][0])
    r.font.size = Pt(12)
    para(f"ORCID {cit['orcid'][0].split('/')[-1]}", size=10, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

    para("  ·  ".join(i["value"] for i in cit["identifiers"]), size=9, color=GREY,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    para(CFG["stamp"].format(sha=git_sha(), date=cit["date_released"]),
         size=8.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

    # ------------------------------------------------------------------ abstract ----- #
    abstract = (SRC / f"sec_abstract{suffix}.md")
    if abstract.exists():
        para(CFG["abstract"], size=13, bold=True, color=NAVY, after=6)
        for block in abstract.read_text().split("\n\n"):
            block = block.strip()
            if not block:
                continue
            if block.startswith("**") and block.endswith("**"):
                para(block.strip("*"), size=10.5, bold=True, after=2)
            else:
                para(block, size=10.5, after=8)

    keywords = (SRC / f"sec_keywords{suffix}.md")
    if keywords.exists():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(CFG["keywords"])
        r.bold = True
        r.font.size = Pt(10)
        _stamp(r, CFG["body"])
        r = p.add_run(keywords.read_text().strip())
        r.font.size = Pt(10)
        _stamp(r, CFG["body"])

    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------------------------ body -------- #
    for b in paper.blocks:
        if b["t"] == "part":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(8)
            r = p.add_run(b["title"])
            r.bold = True
            r.font.size = Pt(15)
            _stamp(r, CFG["sans"])
            r.font.color.rgb = NAVY

        elif b["t"] == "h":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12 if b["level"] <= 2 else 9)
            p.paragraph_format.space_after = Pt(4)
            runs_into(p, b["runs"], size=12.5 if b["level"] <= 2 else 11.5, bold=True,
                      color=NAVY)
            for r in p.runs:
                _stamp(r, CFG["sans"])

        elif b["t"] == "p":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(7)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            runs_into(p, b["runs"])

        elif b["t"] == "li":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            runs_into(p, b["runs"], size=10.5)

        elif b["t"] == "table":
            tb = doc.add_table(rows=1, cols=len(b["head"]))
            tb.style = "Light Grid Accent 1"
            for i, cell in enumerate(b["head"]):
                c = tb.rows[0].cells[i]
                c.text = ""
                runs_into(c.paragraphs[0], cell, size=9, bold=True)
            for row in b["rows"]:
                cells = tb.add_row().cells
                for i, cell in enumerate(row[: len(b["head"])]):
                    cells[i].text = ""
                    runs_into(cells[i].paragraphs[0], cell, size=9)
            doc.add_paragraph().paragraph_format.space_after = Pt(8)

        elif b["t"] == "figure":
            path = FIG / b["name"]
            if not path.exists():
                raise FileNotFoundError(f"{path} is missing; the paper calls for it")
            doc.add_picture(str(path), width=Inches(6.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_after = Pt(14)
            r = cap.add_run(CFG["figure"].format(n=b["n"]))
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = GREY
            runs_into(cap, b["caption"], size=9, color=GREY, italic=True)

    # ------------------------------------------------------------------ back matter - #
    doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(CFG["availability"])
    r.bold = True
    r.font.size = Pt(15)
    _stamp(r, CFG["sans"])
    r.font.color.rgb = NAVY
    para(f"Every prediction, artefact and analysis script is public at "
         f"github.com/hopejsh/CognitionBioChem under Apache-2.0 for the code, with "
         f"third-party scientific data redistributed under CC BY 4.0, CC BY-SA 3.0, CC0 and "
         f"MIT as recorded in NOTICE. {D['runs']} prediction runs are held under "
         f"content-addressed custody and {D['plans']} hash-locked pre-registration plans are "
         f"retained, superseded versions included. Archived at "
         f"{cit['identifiers'][0]['value']}; the resource is registered as "
         f"{cit['identifiers'][2]['value']}.", size=10)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(CFG["references"])
    r.bold = True
    r.font.size = Pt(15)
    _stamp(r, CFG["sans"])
    r.font.color.rgb = NAVY
    para(CFG["refs_note"], size=8.5, italic=True, color=GREY, after=10)

    for r_ in paper.numbered_references():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Inches(0.36)
        p.paragraph_format.first_line_indent = Inches(-0.36)
        num = p.add_run(f"{r_['n']}.  ")
        num.bold = True
        num.font.size = Pt(9)
        _stamp(num, CFG["body"])
        body = p.add_run(render_reference(r_))
        body.font.size = Pt(9)
        _stamp(body, "Cambria")   # a citation stays in a Latin face in both editions
        if r_.get("kind") == "preprint":
            note = p.add_run(CFG["preprint"])
            note.font.size = Pt(8)
            note.italic = True
            note.font.color.rgb = GREY

    doc.save(OUT)
    print(f"wrote {OUT.relative_to(REPO)}")
    print(f"  sections   : {len(ORDER)}")
    print(f"  blocks     : {len(paper.blocks)}")
    print(f"  figures    : {len(paper.figures)}")
    print(f"  references : {len(paper.cited)} cited of {len(paper.refs)} verified")
    words = sum(len(plain(b.get('runs', [])).split()) for b in paper.blocks
                if b['t'] in ('p', 'li', 'h'))
    print(f"  body words : ~{words}")
    return 0


if __name__ == "__main__":
    raise SystemExit(build("ko" if "--ko" in sys.argv else "en"))
